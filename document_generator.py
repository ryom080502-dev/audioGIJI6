"""
ドキュメント生成モジュール - Word/PDF出力
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib import colors
import tempfile
import os
import logging
from typing import Dict
from datetime import datetime

logger = logging.getLogger(__name__)

class DocumentGenerator:
    def __init__(self):
        """ドキュメント生成の初期化"""
        # 日本語フォントの設定（ReportLab用）
        self.japanese_font_available = False
        try:
            # システムの日本語フォントを試行
            # (フォントパス, TTCインデックス) のタプル
            font_configs = [
                ("C:\\Windows\\Fonts\\msgothic.ttc", 0),  # MS ゴシック (Windows)
                ("C:\\Windows\\Fonts\\msmincho.ttc", 0),  # MS 明朝 (Windows)
                ("C:\\Windows\\Fonts\\meiryo.ttc", 0),  # メイリオ (Windows)
                ("C:\\Windows\\Fonts\\yugothic.ttf", None),  # 游ゴシック (Windows)
                ("/System/Library/Fonts/ヒラギノ角ゴシック W3.ttc", 0),  # ヒラギノ (macOS)
                ("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf", None),  # Linux
            ]

            for font_path, ttc_index in font_configs:
                if os.path.exists(font_path):
                    try:
                        # TTCファイルの場合はsubfontIndexを指定
                        if ttc_index is not None and font_path.endswith('.ttc'):
                            pdfmetrics.registerFont(TTFont('Japanese', font_path, subfontIndex=ttc_index))
                        else:
                            pdfmetrics.registerFont(TTFont('Japanese', font_path))

                        self.japanese_font_available = True
                        logger.info(f"日本語フォント登録成功: {font_path}" +
                                  (f" (index: {ttc_index})" if ttc_index is not None else ""))
                        break
                    except Exception as e:
                        logger.debug(f"フォント登録失敗: {font_path} - {str(e)}")
                        continue

            if not self.japanese_font_available:
                logger.warning("日本語フォントが見つかりません。PDF生成時にフォールバックを使用します")

        except Exception as e:
            logger.warning(f"日本語フォント設定エラー: {str(e)}")
    
    def generate_word(self, content: str, metadata: Dict) -> str:
        """
        Word文書を生成
        
        Args:
            content: 議事録の本文
            metadata: メタデータ（日付、作成者など）
            
        Returns:
            生成されたWordファイルのパス
        """
        try:
            logger.info("Word文書の生成を開始")
            
            # 新しい文書を作成
            doc = Document()
            
            # タイトル
            title = doc.add_heading('議事録', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # メタデータテーブル
            doc.add_paragraph()
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Light Grid Accent 1'
            
            # テーブルの内容
            meta_items = [
                ('作成日', metadata.get('created_date', '')),
                ('作成者', metadata.get('creator', '')),
                ('お客様名', metadata.get('customer_name', '')),
                ('打合せ場所', metadata.get('meeting_place', ''))
            ]
            
            for i, (label, value) in enumerate(meta_items):
                row = table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = value
                
                # ラベルセルを太字に
                row.cells[0].paragraphs[0].runs[0].font.bold = True
            
            # 本文
            doc.add_paragraph()
            doc.add_heading('内容', level=1)
            
            # 内容を行ごとに処理
            for line in content.split('\n'):
                line = line.strip()
                if not line:
                    continue

                # セクションヘッダーの判定（##で始まる）
                if line.startswith('##'):
                    heading_text = line.replace('##', '').strip()
                    doc.add_heading(heading_text, level=2)
                # 箇条書きの判定（・、•、-、* で始まる）
                elif line.startswith(('・', '• ', '- ', '* ')):
                    # 箇条書き記号を除去
                    for prefix in ['・', '• ', '- ', '* ']:
                        if line.startswith(prefix):
                            line = line[len(prefix):].strip()
                            break
                    p = doc.add_paragraph(line, style='List Bullet')
                else:
                    # 通常の段落
                    doc.add_paragraph(line)
            
            # フッター
            doc.add_paragraph()
            footer = doc.add_paragraph(f"作成日時: {datetime.now().strftime('%Y年%m月%d日 %H:%M')}")
            footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            footer_run = footer.runs[0]
            footer_run.font.size = Pt(9)
            footer_run.font.color.rgb = RGBColor(128, 128, 128)
            
            # 一時ファイルに保存
            output_path = tempfile.mktemp(suffix=".docx")
            doc.save(output_path)
            
            logger.info(f"Word文書生成完了: {output_path}")
            return output_path
        
        except Exception as e:
            logger.error(f"Word文書生成エラー: {str(e)}")
            raise
    
    def generate_pdf(self, content: str, metadata: Dict) -> str:
        """
        PDF文書を生成
        
        Args:
            content: 議事録の本文
            metadata: メタデータ（日付、作成者など）
            
        Returns:
            生成されたPDFファイルのパス
        """
        try:
            logger.info("PDF文書の生成を開始")
            
            # 一時ファイルパス
            output_path = tempfile.mktemp(suffix=".pdf")
            
            # PDFドキュメント作成
            doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                rightMargin=20*mm,
                leftMargin=20*mm,
                topMargin=20*mm,
                bottomMargin=20*mm
            )
            
            # ストーリー（コンテンツ）のリスト
            story = []
            
            # スタイルの取得と設定
            styles = getSampleStyleSheet()

            # フォント名の設定
            font_name = 'Japanese' if self.japanese_font_available else 'Helvetica'

            # カスタムスタイルの追加
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontName=font_name,
                fontSize=20,
                alignment=TA_CENTER,
                spaceAfter=30,
                textColor=colors.HexColor('#1a1a1a')
            )

            heading1_style = ParagraphStyle(
                'CustomHeading1',
                parent=styles['Heading2'],
                fontName=font_name,
                fontSize=15,
                spaceAfter=15,
                spaceBefore=20,
                textColor=colors.HexColor('#2c3e50'),
                leftIndent=0
            )

            heading2_style = ParagraphStyle(
                'CustomHeading2',
                parent=styles['Heading2'],
                fontName=font_name,
                fontSize=12,
                spaceAfter=10,
                spaceBefore=10,
                textColor=colors.HexColor('#34495e'),
                leftIndent=15
            )

            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['BodyText'],
                fontName=font_name,
                fontSize=10,
                spaceAfter=8,
                leading=16,
                leftIndent=0
            )

            list_style = ParagraphStyle(
                'CustomList',
                parent=styles['BodyText'],
                fontName=font_name,
                fontSize=10,
                spaceAfter=6,
                leading=16,
                leftIndent=25,
                bulletIndent=10
            )

            metadata_style = ParagraphStyle(
                'MetadataStyle',
                parent=styles['BodyText'],
                fontName=font_name,
                fontSize=9,
                spaceAfter=4,
                leading=13,
                textColor=colors.HexColor('#555555')
            )
            
            # タイトル
            story.append(Paragraph('議事録', title_style))
            story.append(Spacer(1, 15))

            # メタデータ
            meta_items = [
                f"<b>作成日:</b> {metadata.get('created_date', '')}",
                f"<b>作成者:</b> {metadata.get('creator', '')}",
                f"<b>お客様名:</b> {metadata.get('customer_name', '')}",
                f"<b>打合せ場所:</b> {metadata.get('meeting_place', '')}"
            ]

            for item in meta_items:
                story.append(Paragraph(item, metadata_style))

            story.append(Spacer(1, 25))

            # 本文を処理
            lines = content.split('\n')
            i = 0
            while i < len(lines):
                line = lines[i].strip()

                if not line:
                    i += 1
                    continue

                # ## で始まる行は大見出し（レベル1）
                if line.startswith('##'):
                    heading_text = line.replace('##', '').strip()
                    story.append(Spacer(1, 5))
                    story.append(Paragraph(heading_text, heading1_style))

                # ・で始まる行は箇条書き
                elif line.startswith('・'):
                    story.append(Paragraph(line, list_style))

                # その他の箇条書き記号
                elif line.startswith(('• ', '- ', '* ')):
                    story.append(Paragraph(line, list_style))

                # 通常のテキスト
                else:
                    # HTMLタグをエスケープ
                    safe_text = line.replace('<', '&lt;').replace('>', '&gt;')

                    # 改行を保持しつつ、適切なスペーシング
                    if safe_text:
                        story.append(Paragraph(safe_text, body_style))

                i += 1
            
            # フッター
            story.append(Spacer(1, 30))
            footer_text = f"作成日時: {datetime.now().strftime('%Y年%m月%d日 %H:%M')}"
            footer_style = ParagraphStyle(
                'Footer',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=9,
                textColor='gray',
                alignment=TA_CENTER
            )
            story.append(Paragraph(footer_text, footer_style))
            
            # PDFビルド
            doc.build(story)
            
            logger.info(f"PDF文書生成完了: {output_path}")
            return output_path
        
        except Exception as e:
            logger.error(f"PDF文書生成エラー: {str(e)}")
            raise
